# app.py
import os
import streamlit as st
import openai
from fpdf import FPDF
from docx import Document
from io import BytesIO
from datetime import datetime

# ---------- CONFIG ----------
# Streamlit Cloud: put OPENAI_API_KEY in Secrets (OPENAI_API_KEY)
OPENAI_KEY = None
if "OPENAI_API_KEY" in st.secrets:
    OPENAI_KEY = st.secrets["OPENAI_API_KEY"]
else:
    OPENAI_KEY = os.getenv("OPENAI_API_KEY")

if not OPENAI_KEY:
    st.warning("OpenAI API key not found. Set OPENAI_API_KEY in Streamlit secrets or environment variable.")
openai.api_key = OPENAI_KEY

MODEL = "gpt-4o"  # change to what you have access to, or "gpt-3.5-turbo"

# ---------- HELPERS ----------
def make_prompt(data):
    # Clear, structured instruction for high-quality resume + cover letter
    return f"""
You are an expert resume writer and career coach. Produce TWO blocks separated with exact tags:
---RESUME---
Format a concise, ATS-friendly resume in plain text (use bullets). Start with Name and Target Title, contact line (email/location if provided), then a 2-3 line Professional Summary tailored to the target role. Then Key Skills (comma/short list). Then Work Experience (reverse chronological: Company — Title — Dates, 3-6 bullets each focusing on achievements and metrics). Then Education. Keep it one page for <10 years experience, two pages max otherwise. Use strong action verbs and metrics when possible.

---COVER LETTER---
Write a cover letter tailored to the role (3 short paragraphs): 1) Open referencing the role and why they're excited, 2) Connect 1-2 concrete achievements/skills to the role, 3) Closing with enthusiasm and call-to-action. Keep professional and concise.

User info:
Name: {data['name']}
Target role: {data['role']}
Location: {data.get('location','')}
Email: {data.get('email','')}
Experience (raw text): {data.get('experience','')}
Education: {data.get('education','')}
Skills: {data.get('skills','')}
Tone: {data.get('tone','professional')}
Extra notes: {data.get('extra','')}

IMPORTANT: Put literal tags ---RESUME--- and ---COVER LETTER---. Output only these two blocks and nothing else.
"""

def call_openai(prompt_text):
    try:
        resp = openai.ChatCompletion.create(
            model=MODEL,
            messages=[
                {"role":"system","content":"You are a professional resume writer."},
                {"role":"user","content":prompt_text}
            ],
            temperature=0.2,
            max_tokens=1600
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"OpenAI error: {e}")
        return None

def text_to_pdf_bytes(title, body_text):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 8, title, ln=True)
    pdf.ln(3)
    pdf.set_font("Arial", size=11)
    for line in body_text.splitlines():
        pdf.multi_cell(0, 6, line)
    out = BytesIO()
    pdf.output(out)
    return out.getvalue()

def text_to_docx_bytes(title, body_text):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in body_text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

# ---------- UI ----------
st.set_page_config(page_title="Strawberry CV — AI Resume", layout="centered", page_icon="🍓")

# top bar with optional logo
col1, col2 = st.columns([1,8])
with col1:
    if os.path.exists("logo.png"):
        st.image("logo.png", width=64)
    else:
        st.markdown("🍓")
with col2:
    st.markdown("<h1 style='margin:0 0 6px 0'>Strawberry CV</h1>", unsafe_allow_html=True)
    st.markdown("<div style='color:gray;margin-top:-10px'>AI-powered resume & cover letter generator</div>", unsafe_allow_html=True)

st.write("---")

with st.form("cv_form"):
    name = st.text_input("Full name", value="")
    role = st.text_input("Target role / Job title", value="")
    email = st.text_input("Contact email (optional)", value="")
    location = st.text_input("Location (optional)", value="")
    experience = st.text_area("Work experience (brief bullets or paragraphs)", value="")
    education = st.text_area("Education (school, degree, dates)", value="")
    skills = st.text_input("Key skills (comma separated)", value="")
    extra = st.text_area("Extra notes (optional) — what to emphasize", value="")
    tone = st.selectbox("Tone", ["professional","friendly","confident","concise"], index=0)
    out_format = st.selectbox("Download formats", ["TXT", "PDF", "DOCX"], index=1)
    submit = st.form_submit_button("Generate Resume & Cover Letter")

if submit:
    if not OPENAI_KEY:
        st.error("No OpenAI API key found. Add OPENAI_API_KEY to Streamlit secrets or environment.")
    else:
        st.info("Generating — please wait...")
        payload = {
            "name": name.strip() or "Candidate Name",
            "role": role.strip() or "Target Role",
            "email": email.strip(),
            "location": location.strip(),
            "experience": experience.strip(),
            "education": education.strip(),
            "skills": skills.strip(),
            "extra": extra.strip(),
            "tone": tone
        }
        prompt = make_prompt(payload)
        result = call_openai(prompt)
        if result:
            resume_text = ""
            cover_text = ""
            if "---RESUME---" in result and "---COVER LETTER---" in result:
                r_idx = result.index("---RESUME---") + len("---RESUME---")
                c_idx = result.index("---COVER LETTER---") + len("---COVER LETTER---")
                resume_text = result[r_idx: result.index("---COVER LETTER---")].strip()
                cover_text = result[c_idx:].strip()
            else:
                # fallback: try split by lines (best-effort)
                lines = result.splitlines()
                resume_text = "\n".join(lines[:40])
                cover_text = "\n".join(lines[40:])

            st.subheader("Generated Resume")
            st.code(resume_text, language=None)

            st.subheader("Generated Cover Letter")
            st.code(cover_text if cover_text else "(no cover letter detected)")

            # prepare downloads
            ts = datetime.now().strftime("%Y%m%d_%H%M")
            base_name = (name or "candidate").replace(" ", "_")
            if out_format == "TXT":
                st.download_button("Download Resume (.txt)", resume_text, file_name=f"{base_name}_resume_{ts}.txt", mime="text/plain")
                st.download_button("Download Cover Letter (.txt)", cover_text, file_name=f"{base_name}_cover_{ts}.txt", mime="text/plain")
            elif out_format == "PDF":
                resume_pdf = text_to_pdf_bytes(f"{name} — Resume", resume_text)
                cover_pdf = text_to_pdf_bytes(f"{name} — Cover Letter", cover_text)
                st.download_button("Download Resume (.pdf)", resume_pdf, file_name=f"{base_name}_resume_{ts}.pdf", mime="application/pdf")
                st.download_button("Download Cover Letter (.pdf)", cover_pdf, file_name=f"{base_name}_cover_{ts}.pdf", mime="application/pdf")
            else:  # DOCX
                resume_docx = text_to_docx_bytes(f"{name} — Resume", resume_text)
                cover_docx = text_to_docx_bytes(f"{name} — Cover Letter", cover_text)
                st.download_button("Download Resume (.docx)", resume_docx, file_name=f"{base_name}_resume_{ts}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.download_button("Download Cover Letter (.docx)", cover_docx, file_name=f"{base_name}_cover_{ts}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            st.success("Generated — review and tweak before sending to employers.")
        else:
            st.error("Generation failed. Check API key, model availability, and quota.")

